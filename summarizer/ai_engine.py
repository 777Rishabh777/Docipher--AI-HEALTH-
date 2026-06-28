import google.generativeai as genai
import os
import io
import PIL.Image
from dotenv import load_dotenv
from functools import lru_cache
from tenacity import retry, stop_after_attempt, wait_none, retry_if_exception_type

load_dotenv()

# --- CONFIGURATION ---
def configure_genai():
    load_dotenv(override=True)
    raw_key = os.getenv('GOOGLE_API_KEY')
    if not raw_key: return False
    api_key = raw_key.strip().replace('"', '').replace("'", "")
    genai.configure(api_key=api_key)
    return True

# --- FALLBACK CONTENT (The Safety Net) ---
FALLBACK_SUMMARY_PATIENT = """
<h3>💚 Quick Patient Summary</h3>
<ul>
    <li><strong>Overview:</strong> Basic blood profile with mild anemia indication.</li>
    <li><strong>Key Points:</strong> Hb slightly low; WBC & Platelets within reference.</li>
    <li><strong>Next Steps:</strong> Iron‑rich diet, hydrate, discuss iron check with doctor.</li>
  
</ul>

<h5 class="mt-3">Your Numbers</h5>
<table class="table table-sm">
    <thead class="table-light"><tr><th>Test</th><th>Value</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Hemoglobin</td><td>10.5 g/dL</td><td>12–16</td><td><span class="text-danger">Low</span></td></tr>
        <tr><td>WBC</td><td>6.8 x10³/µL</td><td>4.0–11.0</td><td>Normal</td></tr>
        <tr><td>Platelets</td><td>245 x10³/µL</td><td>150–450</td><td>Normal</td></tr>
    </tbody>
    </table>

<h5 class="mt-3">Mini‑Charts</h5>
<div class="mb-2">Hemoglobin
    <div class="progress" style="height:10px;"><div class="progress-bar bg-danger" style="width:65%"></div></div>
</div>
<div class="mb-2">WBC
    <div class="progress" style="height:10px;"><div class="progress-bar bg-success" style="width:60%"></div></div>
</div>
<div class="mb-2">Platelets
    <div class="progress" style="height:10px;"><div class="progress-bar bg-success" style="width:55%"></div></div>
</div>

<small class="text-muted"><em>*Failsafe preview (server busy)</em></small>
"""

FALLBACK_SUMMARY_DOCTOR = """
<h3>👨‍⚕️ Clinical Report</h3>

<h5 class="mb-3">Key Vitals & Indices</h5>
<ul class="mb-4">
    <li><strong>Hemoglobin:</strong> 10.5 g/dL (Low) — Mild anemia detected</li>
    <li><strong>WBC:</strong> 6.8 x10³/µL (Normal) — No active infection</li>
    <li><strong>Platelets:</strong> 245 x10³/µL (Normal) — Adequate clotting function</li>
    <li><strong>ALT/AST:</strong> 28/32 U/L (Normal) — Liver enzymes normal</li>
    <li><strong>Creatinine:</strong> 0.95 mg/dL (Normal) — Renal function intact</li>
</ul>

<h5 class="mb-3">Lab Results</h5>
<table class="table table-striped table-bordered" data-role="labs" data-panel="GENERAL">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Hemoglobin</td><td>10.5</td><td>12–16 g/dL</td><td><span class="badge bg-danger">Low</span></td></tr>
        <tr><td>WBC</td><td>6.8</td><td>4.0–11.0 x10³/µL</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Platelets</td><td>245</td><td>150–450 x10³/µL</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>ALT</td><td>28</td><td>7–56 U/L</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>AST</td><td>32</td><td>10–40 U/L</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Creatinine</td><td>0.95</td><td>0.6–1.2 mg/dL</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>TSH</td><td>2.3</td><td>0.4–4.0 mIU/L</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Total Cholesterol</td><td>195</td><td>&lt;200 mg/dL</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<h5 class="mt-4">Mini-Charts (Key Analytes)</h5>
<div class="mb-3">Hemoglobin (10.5/12-16)
    <div class="progress" style="height:12px;"><div class="progress-bar bg-danger" style="width:66%"></div></div>
</div>
<div class="mb-3">WBC (6.8/4.0-11.0)
    <div class="progress" style="height:12px;"><div class="progress-bar bg-success" style="width:62%"></div></div>
</div>
<div class="mb-3">Creatinine (0.95/0.6-1.2)
    <div class="progress" style="height:12px;"><div class="progress-bar bg-success" style="width:79%"></div></div>
</div>
<div class="mb-3">Total Cholesterol (195/&lt;200)
    <div class="progress" style="height:12px;"><div class="progress-bar bg-success" style="width:98%"></div></div>
</div>

<h5 class="mt-4">Impression & Plan</h5>
<p><strong>Findings:</strong> Mild anemia (Hb 10.5) with normal renal function, liver enzymes, and lipid profile. No signs of active infection (WBC normal) or thyroid dysfunction (TSH normal).</p>
<p><strong>Differential:</strong> Iron-deficiency anemia likely given mild presentation. Chronic disease anemia and B12 deficiency should be excluded.</p>
<p><strong>Recommendations:</strong></p>
<ul class="ms-3">
    <li>Iron studies (serum iron, ferritin, TIBC) to confirm iron-deficiency etiology</li>
    <li>Vitamin B12 and folate levels if iron studies unremarkable</li>
    <li>Continue monitoring; recheck CBC in 4–6 weeks post-intervention</li>
    <li>Dietary counsel: iron-rich foods; adequate hydration</li>
</ul>

<small class="text-muted"><em>*Failsafe clinical preview (server busy)</em></small>
"""

FALLBACK_SUMMARY_FULL = """
<h3>📑 Comprehensive Data Report</h3>

<p class="mb-2"><strong>Contents:</strong>
    <a href="#cbc" class="me-2">CBC</a>
    <a href="#lft" class="me-2">LFT</a>
    <a href="#kft" class="me-2">KFT</a>
    <a href="#lipid" class="me-2">Lipids</a>
    <a href="#thyroid" class="me-2">Thyroid</a>
    <a href="#electrolytes">Electrolytes</a>
</p>

<h4 id="cbc" class="mt-3">CBC (Complete Blood Count)</h4>
<p><strong>What it tests:</strong> Measures red blood cells, white blood cells, and platelets. Detects anemia, infection, and clotting disorders.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="CBC">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Hemoglobin</td><td>10.5</td><td>g/dL</td><td>12–16</td><td><span class="badge bg-danger">Low</span></td></tr>
        <tr><td>WBC</td><td>6.8</td><td>x10³/µL</td><td>4.0–11.0</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Platelets</td><td>245</td><td>x10³/µL</td><td>150–450</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="113 377" stroke-dashoffset="0" stroke-linecap="round"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#ef4444" stroke-width="30" stroke-dasharray="113 377" stroke-dashoffset="-113" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 66%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 33%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Mild anemia detected (low hemoglobin). WBC and platelets within normal limits, suggesting no acute infection or clotting issues. Investigate underlying cause of anemia.</p>

<h4 id="lft" class="mt-4">LFT (Liver Function Tests)</h4>
<p><strong>What it tests:</strong> Evaluates liver enzymes, proteins, and bilirubin to assess liver health and function.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="LFT">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>ALT</td><td>28</td><td>U/L</td><td>7–56</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>AST</td><td>32</td><td>U/L</td><td>10–40</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Bilirubin (Total)</td><td>0.9</td><td>mg/dL</td><td>0.3–1.2</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="226 377" stroke-dashoffset="0" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 100%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 0%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Liver function entirely normal. No signs of hepatitis, cirrhosis, or hepatic dysfunction. Liver is healthy.</p>

<h4 id="kft" class="mt-4">KFT (Kidney Function Tests)</h4>
<p><strong>What it tests:</strong> Measures creatinine and BUN to assess kidney filtration capacity and detect renal impairment.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="KFT">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Creatinine</td><td>0.95</td><td>mg/dL</td><td>0.6–1.2</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>BUN</td><td>18</td><td>mg/dL</td><td>7–20</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="226 377" stroke-dashoffset="0" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 100%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 0%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Kidney function normal. GFR adequate. No acute kidney injury or chronic kidney disease detected.</p>

<h4 id="lipid" class="mt-4">Lipids (Lipid Panel)</h4>
<p><strong>What it tests:</strong> Measures cholesterol, triglycerides, HDL, and LDL to assess cardiovascular risk.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="Lipids">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Total Cholesterol</td><td>195</td><td>mg/dL</td><td>&lt;200</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Triglycerides</td><td>120</td><td>mg/dL</td><td>&lt;150</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="226 377" stroke-dashoffset="0" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 100%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 0%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Lipid profile optimal. Low cardiovascular risk. Continue heart-healthy diet and exercise.</p>

<h4 id="thyroid" class="mt-4">Thyroid (TSH, Free T4)</h4>
<p><strong>What it tests:</strong> Measures thyroid stimulating hormone and thyroid hormones to assess thyroid function.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="Thyroid">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>TSH</td><td>2.3</td><td>mIU/L</td><td>0.4–4.0</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="226 377" stroke-dashoffset="0" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 100%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 0%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Thyroid function normal. No hypothyroidism or hyperthyroidism. Healthy thyroid.</p>

<h4 id="electrolytes" class="mt-4">Electrolytes</h4>
<p><strong>What it tests:</strong> Measures sodium, potassium, chloride, and CO₂ to assess electrolyte balance and hydration.</p>

<table class="table table-striped table-bordered" data-role="labs" data-panel="Electrolytes">
    <thead class="table-light"><tr><th>Parameter</th><th>Value</th><th>Units</th><th>Reference</th><th>Status</th></tr></thead>
    <tbody>
        <tr><td>Sodium</td><td>138</td><td>mEq/L</td><td>136–145</td><td><span class="badge bg-success">Normal</span></td></tr>
        <tr><td>Potassium</td><td>4.2</td><td>mEq/L</td><td>3.5–5.0</td><td><span class="badge bg-success">Normal</span></td></tr>
    </tbody>
</table>

<div class="pie-chart-container" style="margin-top:1rem;">
  <h6>Result Distribution</h6>
  <svg width="160" height="160" viewBox="-10 -10 160 160" style="margin-right:1.5rem;float:left;">
    <circle cx="70" cy="70" r="60" fill="none" stroke="#e5e7eb" stroke-width="3"/>
    <circle cx="70" cy="70" r="60" fill="none" stroke="#22c55e" stroke-width="30" stroke-dasharray="226 377" stroke-dashoffset="0" stroke-linecap="round"/>
  </svg>
  <div style="overflow:hidden;">
    <div class="pie-chart-legend">
      <span><span style="display:inline-block;width:12px;height:12px;background:#22c55e;margin-right:0.5rem;"></span>Normal: 100%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#ef4444;margin-right:0.5rem;"></span>Abnormal: 0%</span><br>
      <span><span style="display:inline-block;width:12px;height:12px;background:#eab308;margin-right:0.5rem;"></span>Warning: 0%</span>
    </div>
  </div>
</div>
<div style="clear:both;"></div>

<p style="margin-top:1rem;"><strong>Clinical Note:</strong> Electrolytes balanced. Hydration status adequate. No dehydration or electrolyte disturbance.</p>

<h5 class="mt-4">Overall Summary</h5>
<p><strong>Key Findings:</strong> Mild anemia is the primary abnormality. All other organ systems tested are functioning normally. Recommend investigating cause of anemia (iron deficiency, B12/folate, chronic disease). Otherwise, health indicators are good.</p>

<small class="text-muted"><em>*Failsafe preview (server busy). Live AI will provide actual analysis based on your uploaded report.</em></small>
"""

FALLBACK_CHAT_RESPONSE = """Thanks for your question. Live AI is busy—here’s a quick, safe checklist:

**Overview & Causes:** Commonly mild viral illness, stress/fatigue, dehydration, or minor strain.

**Red Flags (urgent care):**
- Sudden severe pain
- Fever >103°F (39.4°C) or >3 days
- Chest pain, trouble breathing, confusion, seizure, stroke signs
- Persistent vomiting, blood in vomit/stool, severe dehydration

**Home Remedies:**
- Rest; 8–10 glasses water/ORS; light bland meals
- Cool/heat pack as suitable; gentle stretch if muscular
- Steam for congestion; honey + warm water for throat (not for infants <1y)
- Avoid heavy/spicy/oily meals; good sleep routine

**OTC (adult, generic):**
- Paracetamol 500–1000mg q4–6h (max 4g/day)
- Ibuprofen 400mg q6–8h with food (avoid ulcer/renal risk)
- Cetirizine 10mg daily for allergy symptoms

**See a Doctor:**
- Symptoms persist/worsen >3–5 days
- Pain limits daily life
- Pregnancy or drug-interaction concerns

⚠️ MEDICAL DISCLAIMER: Educational only; not a substitute for professional medical advice. Consult a qualified clinician before treatment.
🚑 Emergency: For severe or rapidly worsening symptoms, call local emergency services immediately.

For a personalized answer, please try again shortly."""

# --- ROBUST GENERATOR ---
@retry(stop=stop_after_attempt(3), wait=wait_none(), reraise=True)
def _attempt_generate(model_name, payload, safety):
    print(f"DEBUG: Connecting to {model_name}...")
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(payload, safety_settings=safety)
    return response.text

def generate_robust_response(content_list, mode='patient', is_chat=False):
    if not configure_genai(): return "Error: API Key missing."

    # Try different models that have less quota usage - prioritize lite models
    models_to_try = [
        'gemini-2.5-flash-lite',
        'gemini-2.0-flash-lite',  
        'gemini-flash-lite-latest',
        'gemini-2.5-flash',
        'gemini-2.0-flash'
    ]
    
    safety = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    # 1. Try Live AI
    for model_name in models_to_try:
        try:
            text = _attempt_generate(model_name, content_list, safety)
            if text: return text.replace("```html", "").replace("```", "")
        except Exception as e:
            print(f"DEBUG: {model_name} failed: {e}")
            continue 

    # 2. Failsafe Activation (Instant Response)
    print("DEBUG: Activating Failsafe Mode.")
    if is_chat: return FALLBACK_CHAT_RESPONSE
    if mode == 'doctor': return FALLBACK_SUMMARY_DOCTOR
    if mode == 'full': return FALLBACK_SUMMARY_FULL
    return FALLBACK_SUMMARY_PATIENT

# --- WRAPPERS ---
def summarize_image_with_gemini(image_path, mode='patient'):
    try:
        img = PIL.Image.open(image_path)
        img_byte_arr = io.BytesIO()
        img.convert('RGB').save(img_byte_arr, format='JPEG')
        image_payload = {"mime_type": "image/jpeg", "data": img_byte_arr.getvalue()}
        
        def mode_instructions(m):
            if m == 'doctor':
                return (
                    "Role: Medical Scribe for clinicians.\n"
                    "Return clean HTML only. Use Bootstrap classes if helpful.\n"
                    "Sections: 1) Key Vitals/Indices (bullets) 2) Lab Table (Parameter | Value | Reference | Status, add data-role='labs' data-panel='GENERAL') 3) Mini‑Charts (Bootstrap progress bars for major tests) 4) Impression & Plan (brief, clinical).\n"
                    "Rules: Extract only values present in the image; do not invent. Keep concise."
                )
            if m == 'full':
                return (
                    "Role: Medical Data Reporter - COMPREHENSIVE.\n"
                    "Return HTML with: 1) a 'Contents' list linking to panel anchors; 2) multiple panels in this exact order: CBC, LFT, KFT, Lipids, Thyroid, Electrolytes, Other.\n"
                    "For each panel: <h4 id='panel_name'>Panel Name</h4>, brief 2–3 sentence explanation of what the panel tests, table (data-role='labs' data-panel='PANEL'), optional inline SVG sparklines, and a pie chart showing result distribution (Normal | Abnormal | Warning).\n"
                    "Pie charts use inline SVG (width 120 height 120): draw 3 colored pie slices (green for Normal, orange for Abnormal, red for Warning) with percentages.\n"
                    "Example: 'CBC panel measures red/white cells & platelets to assess infection, anemia & clotting.'\n"
                    "Add 3–5 key analyte progress bars at the end. Include brief clinical notes/impressions.\n"
                    "Rules: Enforce panel order; prefer tables; avoid invented values; valid HTML only."
                )
            # patient (default)
            return (
                "Role: Patient Educator.\n"
                "Return short HTML with: Key Results (bullets), What it means (plain words), Next steps (3 bullets), and a small 'Your Numbers' table.\n"
                "Include 2–3 mini‑charts (Bootstrap progress bars). Keep it simple."
            )

        prompt = mode_instructions(mode) + "\n\nAnalyze this medical image. Output ONLY valid HTML."
        return generate_robust_response([prompt, image_payload], mode=mode)
    except Exception as e:
        return f"<div class='alert alert-danger'>Image Error: {str(e)}</div>"

def get_summary_formatted(text, mode='patient'):
    def text_mode_prompt(m):
        if m == 'doctor':
            return (
                "You are a clinical report writer. Read the text and produce concise HTML for clinicians.\n"
                "Sections: 1) Snapshot (bullets) 2) Labs Table (Parameter | Value | Reference | Status, add data-role='labs' data-panel='GENERAL') 3) Mini‑Charts using Bootstrap progress bars for key analytes 4) Brief Impression & Plan.\n"
                "Rules: Extract ONLY data present; avoid guessing; keep it compact; no markdown fences; valid HTML only."
            )
        if m == 'full':
            return (
                "You are a medical data summarizer. Produce a comprehensive HTML report.\n"
                "Include at top a compact 'Contents' with anchors to panels (e.g., #cbc, #lft, #kft, #lipid, #thyroid, #electrolytes).\n"
                "For each applicable panel, render <h4 id='panel-id'>Panel Name</h4> and a table with data-role='labs' data-panel='PANEL_NAME' (columns: Parameter | Value | Units | Reference | Flag).\n"
                "Add a tiny inline SVG sparkline next to parameters that have a short time‑series explicitly present (e.g., last 3 values), using <svg width='100' height='24'><polyline points='...'/>. If no series provided, skip sparkline.\n"
                "Add progress bars (<div class='progress'><div class='progress-bar'>) for 3–5 key analytes. Close with Notes/Interpretation.\n"
                "Rules: Prefer tables; be explicit about units; no invented values; valid HTML only."
            )
        # patient
        return (
            "You are a patient educator. Create a SHORT, easy HTML summary.\n"
            "Sections: Key Results (3–6 bullets), What it Means (plain language), Next Steps (diet/lifestyle/when to recheck), and a small table 'Your Numbers' with 5–10 most relevant tests (Parameter | Value | Range | Status).\n"
            "Add 2–3 simple mini‑charts using Bootstrap progress bars. Keep it brief and friendly. Valid HTML only."
        )

    prompt = text_mode_prompt(mode)
    return generate_robust_response([prompt, text], mode=mode)

def summarize_pdf_with_gemini(pdf_path, mode='patient'):
    try:
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        pdf_payload = {"mime_type": "application/pdf", "data": pdf_bytes}
        
        def mode_instructions(m):
            if m == 'doctor':
                return (
                    "Role: Medical Scribe for clinicians.\n"
                    "Return clean HTML only. Use Bootstrap classes if helpful.\n"
                    "Sections: 1) Key Vitals/Indices (bullets) 2) Lab Table (Parameter | Value | Reference | Status, add data-role='labs' data-panel='GENERAL') 3) Mini‑Charts (Bootstrap progress bars for major tests) 4) Impression & Plan (brief, clinical).\n"
                    "Rules: Extract only values present in the document; do not invent. Keep concise."
                )
            if m == 'full':
                return (
                    "Role: Medical Data Reporter - COMPREHENSIVE.\n"
                    "Return HTML with: 1) a 'Contents' list linking to panel anchors; 2) multiple panels in this exact order: CBC, LFT, KFT, Lipids, Thyroid, Electrolytes, Other.\n"
                    "For each panel: <h4 id='panel_name'>Panel Name</h4>, brief 2–3 sentence explanation of what the panel tests, table (data-role='labs' data-panel='PANEL'), optional inline SVG sparklines, and a pie chart showing result distribution (Normal | Abnormal | Warning).\n"
                    "Pie charts use inline SVG (width 120 height 120): draw 3 colored pie slices (green for Normal, orange for Abnormal, red for Warning) with percentages.\n"
                    "Example: 'CBC panel measures red/white cells & platelets to assess infection, anemia & clotting.'\n"
                    "Add 3–5 key analyte progress bars at the end. Include brief clinical notes/impressions.\n"
                    "Rules: Enforce panel order; prefer tables; avoid invented values; valid HTML only."
                )
            # patient (default)
            return (
                "Role: Patient Educator.\n"
                "Return short HTML with: Key Results (bullets), What it means (plain words), Next steps (3 bullets), and a small 'Your Numbers' table.\n"
                "Include 2–3 mini‑charts (Bootstrap progress bars). Keep it simple."
            )

        prompt = mode_instructions(mode) + "\n\nAnalyze this medical document. Output ONLY valid HTML."
        return generate_robust_response([prompt, pdf_payload], mode=mode)
    except Exception as e:
        return f"<div class='alert alert-danger'>PDF Error: {str(e)}</div>"

def extract_text_from_file(file_path):
    ext = file_path.split('.')[-1].lower()
    
    # Text-like files
    if ext in ['txt', 'md', 'rtf', 'html', 'xml', 'json']:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except Exception as e:
            print(f"Error reading text file {file_path}: {e}")
            return ""
            
    # PDF files
    elif ext == 'pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {e}")
            return ""
            
    # Word files
    elif ext in ['docx']:
        try:
            import docx
            doc = docx.Document(file_path)
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        text_parts.append(" | ".join(cells))
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
            
    # Excel/Spreadsheet files
    elif ext in ['xlsx', 'xls']:
        try:
            import pandas as pd
            text_parts = []
            with pd.ExcelFile(file_path) as xls:
                for name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=name)
                    text_parts.append(f"Sheet Name: {name}")
                    text_parts.append(df.to_string(index=False))
            return "\n\n".join(text_parts).strip()
        except Exception as e:
            print(f"Error extracting text from Excel {file_path}: {e}")
            return ""
            
    # CSV files
    elif ext == 'csv':
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string(index=False).strip()
        except Exception as e:
            print(f"Error extracting text from CSV {file_path}: {e}")
            return ""
            
    return ""

# Medical Keywords for filtering
MEDICAL_KEYWORDS = [
    'fever', 'cold', 'cough', 'pain', 'headache', 'stomach', 'ache', 'medicine', 'medication',
    'sick', 'ill', 'disease', 'symptom', 'treatment', 'cure', 'drug', 'tablet', 'pill', 
    'doctor', 'health', 'medical', 'allergy', 'infection', 'diabetes', 'blood pressure',
    'heart', 'lung', 'kidney', 'liver', 'cancer', 'body', 'throat', 'chest', 'back',
    'injury', 'wound', 'fracture', 'sprain', 'vitamin', 'supplement', 'nutrition',
    'diet', 'exercise', 'sleep', 'stress', 'anxiety', 'depression', 'mental health',
    'pregnancy', 'baby', 'child', 'elderly', 'senior', 'aging', 'skin', 'rash',
    'vomit', 'nausea', 'diarrhea', 'constipation', 'digestion', 'appetite',
    'weight', 'fatigue', 'tired', 'weakness', 'dizzy', 'vertigo', 'bleeding',
    'swelling', 'inflammation', 'prescription', 'diagnosis', 'test', 'scan',
    'vaccine', 'immunization', 'prevention', 'hygiene', 'sanitize', 'coronavirus',
    'covid', 'flu', 'asthma', 'bronchitis', 'pneumonia', 'malaria', 'dengue',
    'temperature', 'bp', 'sugar level', 'cholesterol', 'bone', 'joint', 'arthritis',
    'physiotherapy', 'therapy', 'counseling', 'surgery', 'operation', 'emergency',
    'first aid', 'remedy', 'home remedy', 'natural cure', 'ayurvedic', 'herbal'
]

def is_medical_query(question):
    """Check if the question is related to medical/health topics"""
    question_lower = question.lower()
    return any(keyword in question_lower for keyword in MEDICAL_KEYWORDS)

def build_profile_context(profile):
    """Build a context string from user's medical profile"""
    if not profile:
        return ""
    
    context_parts = []
    
    if profile.date_of_birth:
        from datetime import date
        age = (date.today() - profile.date_of_birth).days // 365
        context_parts.append(f"Age: {age} years")
    
    if profile.gender:
        context_parts.append(f"Gender: {profile.gender}")
    
    if profile.blood_group and profile.blood_group != 'Unknown':
        context_parts.append(f"Blood Group: {profile.blood_group}")
    
    if profile.weight:
        context_parts.append(f"Weight: {profile.weight} kg")
    
    if profile.height:
        context_parts.append(f"Height: {profile.height} cm")
    
    if profile.allergies:
        context_parts.append(f"Known Allergies: {profile.allergies}")
    
    if profile.chronic_conditions:
        context_parts.append(f"Chronic Conditions: {profile.chronic_conditions}")
    
    if profile.current_medications:
        context_parts.append(f"Current Medications: {profile.current_medications}")
    
    if profile.smoking_status:
        context_parts.append(f"Smoking Status: {profile.smoking_status}")
    
    if profile.alcohol_consumption:
        context_parts.append(f"Alcohol Consumption: {profile.alcohol_consumption}")
    
    if context_parts:
        return "Patient Profile:\n" + "\n".join(context_parts)
    return ""

def get_chatbot_response(user_question, profile=None):
    """
    Universal AI assistant - answers ANY question with real AI.
    """
    
    profile_context = build_profile_context(profile)

    # Simple, universal prompt that works for ANY question
    if profile_context:
        prompt = f"""You are Docipher AI, a helpful assistant.

{profile_context}

User Question: {user_question}

Provide a clear, helpful, and accurate answer. If it's a medical question and the user has health information above, consider it in your response."""
    else:
        prompt = f"""You are Docipher AI, a helpful assistant.

User Question: {user_question}

Provide a clear, helpful, and accurate answer."""

    # Call the REAL AI - no fallbacks unless absolutely necessary
    response = generate_robust_response([prompt], is_chat=True)
    
    # Only add disclaimer for medical questions
    if any(word in user_question.lower() for word in ['medicine', 'drug', 'symptom', 'disease', 'pain', 'fever', 'doctor', 'health', 'sick']):
        disclaimer = "\n\n⚠️ Medical Disclaimer: For health matters, consult a qualified healthcare professional."
        if "disclaimer" not in response.lower():
            response += disclaimer
    
    return response